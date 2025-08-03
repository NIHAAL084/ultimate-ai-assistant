import pdfplumber
import pytesseract  # type: ignore
import fitz  # type: ignore
from PIL import Image
import io
from typing import Any, Dict, List
from docx import Document
pytesseract.pytesseract.tesseract_cmd = r'/usr/local/bin/tesseract' # For macOS/Linux


def extract_universal_pdf_data(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text and tables from each page of a PDF, automatically using OCR
    for image-based pages.

    Args:
        pdf_path: The file path to the PDF.

    Returns:
        A list of dictionaries, where each dictionary represents a page
        and contains its 'page_number', extracted 'text', and 'tables'.
    """
    all_page_data: List[Dict[str, Any]] = []
    # Use PyMuPDF to handle the rasterization for OCR
    doc_for_ocr = fitz.open(pdf_path)

    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_data: Dict[str, Any] = {
                "page_number": i + 1,
                "text": "",
                "tables": []
            }

            # First, try direct text extraction
            text = page.extract_text()

            # Heuristic: If a page has very little text, it's likely scanned
            if not text or len(text.strip()) < 50:
                # Use PyMuPDF to get an image of the page
                p_ocr = doc_for_ocr.load_page(i)  # type: ignore
                # Render page to an image (300 DPI for better OCR)
                pix = p_ocr.get_pixmap(dpi=300)  # type: ignore
                img_data = pix.tobytes("png")  # type: ignore
                image = Image.open(io.BytesIO(img_data))  # type: ignore
                
                # Perform OCR
                ocr_text = pytesseract.image_to_string(image, lang='eng')  # type: ignore
                page_data["text"] = ocr_text
                # Note: Table extraction is not feasible on OCR'd images directly
                
            else:
                # If text exists, use it and extract tables
                page_data["text"] = text
                # Extract tables with pdfplumber's excellent engine
                page_data["tables"] = page.extract_tables() or []

            all_page_data.append(page_data)  # type: ignore


    doc_for_ocr.close()
    return all_page_data



def extract_docx_data(docx_path: str) -> Dict[str, Any]:
    """
    Extracts all text and tables from a .docx file.

    Args:
        docx_path: The file path to the .docx document.

    Returns:
        A dictionary containing the full text and a list of all tables.
    """
    doc = Document(docx_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)  # type: ignore


    # Extract data from tables
    all_tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)  # type: ignore
        all_tables.append(table_data)  # type: ignore

    return {
        "text": "\n".join(full_text),  # type: ignore
        "tables": all_tables
    }
